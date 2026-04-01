"""
MCP Server - 将 Word 文档中的批注（comments）内联到正文中。
"""

import base64
import io
import os
import tempfile

import lxml.etree as ET
from docx import Document
from mcp.server.fastmcp import FastMCP

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

mcp = FastMCP("docx-comments")


def _extract_comments(doc):
    """从文档中提取所有批注"""
    comments = {}
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype and "Extended" not in rel.reltype:
            comments_xml = ET.fromstring(rel.target_part.blob)
            for comment in comments_xml.findall(f".//{{{W}}}comment"):
                cid = comment.get(f"{{{W}}}id")
                texts = []
                for t in comment.iter(f"{{{W}}}t"):
                    if t.text:
                        texts.append(t.text)
                comments[cid] = "".join(texts)
    return comments


def _make_run(text):
    """创建一个红色加粗的批注 run 元素"""
    run = ET.Element(f"{{{W}}}r")
    rPr = ET.SubElement(run, f"{{{W}}}rPr")
    c = ET.SubElement(rPr, f"{{{W}}}color")
    c.set(f"{{{W}}}val", "FF0000")
    ET.SubElement(rPr, f"{{{W}}}b")
    sz = ET.SubElement(rPr, f"{{{W}}}sz")
    sz.set(f"{{{W}}}val", "21")
    szCs = ET.SubElement(rPr, f"{{{W}}}szCs")
    szCs.set(f"{{{W}}}val", "21")
    t_elem = ET.SubElement(run, f"{{{W}}}t")
    t_elem.set(XML_SPACE, "preserve")
    t_elem.text = f"【批注：{text}】"
    return run


def _inline_comments(doc, comments):
    """将批注插入到正文对应位置"""
    body = doc.element.body
    processed = set()

    for range_end in list(body.iter(f"{{{W}}}commentRangeEnd")):
        cid = range_end.get(f"{{{W}}}id")
        if cid in comments and cid not in processed:
            processed.add(cid)
            parent = range_end.getparent()
            idx = list(parent).index(range_end)
            parent.insert(idx + 1, _make_run(comments[cid]))

    for ref in list(body.iter(f"{{{W}}}commentReference")):
        cid = ref.get(f"{{{W}}}id")
        if cid in comments and cid not in processed:
            processed.add(cid)
            run_parent = ref.getparent()
            p_parent = run_parent.getparent()
            idx = list(p_parent).index(run_parent)
            p_parent.insert(idx + 1, _make_run(comments[cid]))

    return processed


@mcp.tool()
def extract_comments(file_base64: str = "", file_path: str = "") -> str:
    """提取 Word 文档中的所有批注内容。

    支持两种输入方式（二选一）：
    - file_base64: 传入 base64 编码的 .docx 文件内容（适用于通过聊天上传文件）
    - file_path: 传入本地 .docx 文件路径

    Args:
        file_base64: base64 编码的 .docx 文件内容
        file_path: 本地 .docx 文件路径

    Returns:
        文档中所有批注的文本内容
    """
    if file_base64:
        file_bytes = base64.b64decode(file_base64)
        doc = Document(io.BytesIO(file_bytes))
    elif file_path:
        if not os.path.exists(file_path):
            return f"错误：文件不存在 - {file_path}"
        doc = Document(file_path)
    else:
        return "错误：请提供 file_base64 或 file_path 参数"

    comments = _extract_comments(doc)

    if not comments:
        return "该文档中没有找到批注。"

    lines = [f"共找到 {len(comments)} 条批注：\n"]
    for i, (cid, text) in enumerate(comments.items(), 1):
        lines.append(f"{i}. {text}")

    return "\n".join(lines)


@mcp.tool()
def inline_comments_base64(file_base64: str, filename: str = "document.docx") -> str:
    """将 Word 文档中的批注内联到正文中。

    接收 base64 编码的 docx 文件，将批注以【批注：内容】的格式（红色加粗）插入到正文对应位置，
    返回处理后的 base64 编码文件。

    Args:
        file_base64: base64 编码的 .docx 文件内容
        filename: 文件名（可选，用于生成输出文件名）

    Returns:
        JSON 字符串，包含处理后的 base64 文件和统计信息
    """
    import json

    file_bytes = base64.b64decode(file_base64)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
        tmp_in.write(file_bytes)
        tmp_in_path = tmp_in.name

    try:
        doc = Document(tmp_in_path)
        comments = _extract_comments(doc)
        processed = _inline_comments(doc, comments)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_out:
            tmp_out_path = tmp_out.name

        doc.save(tmp_out_path)

        with open(tmp_out_path, "rb") as f:
            result_bytes = f.read()

        os.unlink(tmp_out_path)

        base, ext = os.path.splitext(filename)
        output_filename = f"{base}_批注内联{ext}"

        return json.dumps({
            "filename": output_filename,
            "file_base64": base64.b64encode(result_bytes).decode(),
            "total_comments": len(comments),
            "processed_comments": len(processed),
        }, ensure_ascii=False)
    finally:
        os.unlink(tmp_in_path)


@mcp.tool()
def inline_comments_file(input_path: str, output_path: str = "") -> str:
    """将本地 Word 文档中的批注内联到正文中。

    读取本地 docx 文件，将批注以【批注：内容】的格式（红色加粗）插入到正文对应位置，
    保存处理后的文件。

    Args:
        input_path: 输入的 .docx 文件路径
        output_path: 输出文件路径（可选，默认在原文件名后加 _批注内联）

    Returns:
        处理结果信息
    """
    if not os.path.exists(input_path):
        return f"错误：文件不存在 - {input_path}"

    doc = Document(input_path)
    comments = _extract_comments(doc)
    processed = _inline_comments(doc, comments)

    if not output_path:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_批注内联{ext}"

    doc.save(output_path)

    return (
        f"处理完成！找到 {len(comments)} 条批注，"
        f"已处理 {len(processed)} 条。\n"
        f"输出文件：{output_path}"
    )


def main():
    mcp.run(transport="stdio")


if __name__ == "__main__":
    main()
